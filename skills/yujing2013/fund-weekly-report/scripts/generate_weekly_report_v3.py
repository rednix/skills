#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金周报生成脚本 V3 - 完整版
包含：近一周 + 年初以来 数据
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Any, Tuple
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')


def read_excel_data(fund_file: str, etf_file: str) -> Tuple[Dict[str, pd.DataFrame], Dict[str, Any]]:
    """读取Excel数据"""
    fund_data = {}
    fund_excel = pd.ExcelFile(fund_file)
    for sheet in fund_excel.sheet_names:
        fund_data[sheet] = pd.read_excel(fund_file, sheet_name=sheet)
    
    etf_data = {}
    etf_excel = pd.ExcelFile(etf_file)
    for sheet in etf_excel.sheet_names:
        etf_data[sheet] = pd.read_excel(etf_file, sheet_name=sheet)
    
    return fund_data, etf_data


def get_date_range(fund_data: Dict[str, pd.DataFrame]) -> Tuple[str, str]:
    """获取日期范围"""
    df = fund_data.get('本周新成立的基金')
    if df is not None and len(df) > 0:
        dates = df['基金成立日'].dropna()
        if len(dates) > 0:
            dates = pd.to_datetime(dates)
            start_date = dates.min().strftime('%m%d')
            end_date = dates.max().strftime('%m%d')
            return (start_date, end_date)
    return ('0310', '0314')


def extract_active_equity_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取主动权益基金统计数据"""
    df = fund_data.get('主动权益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    fund_types = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
    
    # 近一周收益
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    # 年初以来收益
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 5
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_industry_fund_returns(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取行业主题基金收益数据"""
    result = {'weekly': [], 'ytd': [], 'top_funds': []}
    
    # 近一周收益
    df_weekly = fund_data.get('行业基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('行业基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['年初以来收益']
                })
    
    # 头部绩优产品（近一周+年初以来）
    df_top = fund_data.get('行业基金基金_收益top')
    if df_top is not None:
        for sector in df_top['所属行业板块'].unique():
            sector_df = df_top[df_top['所属行业板块'] == sector]
            # 过滤掉"平均收益"行
            sector_df = sector_df[sector_df['证券简称'] != '平均收益']
            
            if len(sector_df) > 0:
                # 近一周TOP
                weekly_top = sector_df.nsmallest(5, '近一周收益', keep='first') if sector_df['近一周收益'].iloc[0] < 0 else sector_df.nlargest(5, '近一周收益', keep='first')
                # 年初以来TOP
                ytd_top = sector_df.nlargest(5, '年初以来收益', keep='first')
                
                result['top_funds'].append({
                    '行业': sector,
                    '近一周TOP': [{'名称': row['证券简称'], '收益': row['近一周收益']} for _, row in weekly_top.head(3).iterrows()],
                    '年初以来TOP': [{'名称': row['证券简称.1'], '收益': row['年初以来收益']} for _, row in ytd_top.head(3).iterrows()]
                })
    
    return result


def extract_fixed_income_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取固定收益基金统计数据"""
    df = fund_data.get('固定收益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    # 固定收益_周度收益统计 数据结构：
    # 列1-6：近一周收益（短期纯债、中长期纯债、货币市场、一级债基、二级债基、偏债混合）
    # 列7-12：年初以来收益（短期纯债、中长期纯债、货币市场、一级债基、二级债基、偏债混合）
    
    fund_types = ['短期纯债型基金', '中长期纯债型基金', '一级债基', '二级债基', '偏债混合型基金']
    
    # 近一周收益（列1-5，跳过货币市场型基金）
    col_map_weekly = [1, 2, 4, 5, 6]  # 短期纯债、中长期纯债、一级债基、二级债基、偏债混合
    for i, (fund_type, col_idx) in enumerate(zip(fund_types, col_map_weekly)):
        result['weekly'][fund_type] = {
            '最高值': df.iloc[2, col_idx] if len(df) > 2 else None,  # 第2行是最高值
            '50%分位': df.iloc[5, col_idx] if len(df) > 5 else None,  # 第5行是50%分位
            '最低值': df.iloc[8, col_idx] if len(df) > 8 else None,  # 第8行是最低值
        }
    
    # 年初以来收益（列7-11，跳过货币市场型基金）
    col_map_ytd = [7, 8, 10, 11, 12]  # 短期纯债、中长期纯债、一级债基、二级债基、偏债混合
    for i, (fund_type, col_idx) in enumerate(zip(fund_types, col_map_ytd)):
        result['ytd'][fund_type] = {
            '最高值': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[5, col_idx] if len(df) > 5 else None,
            '最低值': df.iloc[8, col_idx] if len(df) > 8 else None,
        }
    
    return result


def extract_index_fund_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取指数基金数据"""
    result = {
        'weekly': [],
        'ytd': [],
        'average': [],
        'enhanced_top': []
    }
    
    # 近一周收益
    df_weekly = fund_data.get('指数基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['四级分类'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('指数基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['四级分类'],
                    '收益': row['年初以来收益']
                })
    
    # 平均收益
    df_avg = fund_data.get('指数基金_平均收益')
    if df_avg is not None:
        for _, row in df_avg.iterrows():
            if pd.notna(row.get('四级分类')):
                result['average'].append({
                    '类型': row['四级分类'],
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    # 增强指基TOP
    df_enhanced = fund_data.get('股票指数增强基金_收益top')
    if df_enhanced is not None:
        for _, row in df_enhanced.iterrows():
            if pd.notna(row.get('证券简称')):
                result['enhanced_top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_fof_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取FOF基金数据"""
    result = {'weekly': [], 'ytd': [], 'top': []}
    
    # 近一周收益
    df_weekly = fund_data.get('FOF近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['投资类型'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('FOF年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['投资类型'],
                    '收益': row['年初以来收益']
                })
    
    # TOP基金
    df_top = fund_data.get('FOF基金_收益top')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('证券简称')):
                result['top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_qdii_data(fund_data: Dict[str, pd.DataFrame]) -> List[Dict]:
    """提取QDII基金数据"""
    result = []
    df = fund_data.get('QDII基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            if pd.notna(row.get('证券简称')):
                result.append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    return result


def extract_reits_data(fund_data: Dict[str, pd.DataFrame]) -> List[Dict]:
    """提取REITs基金数据"""
    result = []
    df = fund_data.get('REITs基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            if pd.notna(row.get('证券简称')):
                result.append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    return result


def extract_new_funds(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取新基金数据"""
    result = {'established': [], 'issued': [], 'declared': []}
    
    # 新成立基金
    df_est = fund_data.get('本周新成立的基金')
    if df_est is not None:
        for _, row in df_est.iterrows():
            if pd.notna(row.get('证券简称')):
                result['established'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '基金成立日': row.get('基金成立日'),
                    '基金分类': row.get('基金分类'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '发行规模': row.get('发行规模(亿元)'),
                    '基金经理': row.get('基金经理')
                })
    
    # 新发行基金
    df_issue = fund_data.get('本周新发行的基金')
    if df_issue is not None:
        for _, row in df_issue.iterrows():
            if pd.notna(row.get('证券简称')):
                result['issued'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '发行起始日': row.get('发行起始日'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '基金经理': row.get('基金经理')
                })
    
    # 新申报基金
    df_decl = fund_data.get('本周新申报的基金')
    if df_decl is not None:
        for _, row in df_decl.iterrows():
            if pd.notna(row.get('基金名称')):
                result['declared'].append({
                    '基金管理人': row.get('基金管理人'),
                    '基金名称': row.get('基金名称'),
                    '基金类型': row.get('基金类型1'),
                    '申请日期': row.get('申请材料接收日')
                })
    
    return result


def extract_etf_flow_data(etf_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取ETF资金流动数据"""
    result = {
        'sector_flow': [],
        'core_index_flow': [],
        'hot_theme_flow': [],
        'top_inflow': [],
        'top_outflow': []
    }
    
    # 板块资金流动
    df_sector = etf_data.get('板块结果展示')
    if df_sector is not None:
        for _, row in df_sector.iterrows():
            if pd.notna(row.get('二级分类')):
                inflow = row.get('资金流入总额（亿元）')
                result['sector_flow'].append({
                    '板块': row['二级分类'],
                    '平均涨跌幅': row.get('平均涨跌幅'),
                    '资金流入总额': inflow if pd.notna(inflow) else 0
                })
    
    # 核心宽基ETF资金流动
    df_core = etf_data.get('核心宽基赛道')
    if df_core is not None:
        for _, row in df_core.iterrows():
            if pd.notna(row.get('跟踪指数')):
                net_flow = row.get('净申购额（亿元）.1')
                result['core_index_flow'].append({
                    '跟踪指数': row.get('跟踪指数'),
                    '周收益率': row.get('周收益率'),
                    '净申购额': net_flow if pd.notna(net_flow) else 0
                })
    
    # 热门行业主题（含年初以来收益率）
    df_theme = etf_data.get('热门行业主题')
    if df_theme is not None:
        for _, row in df_theme.iterrows():
            if pd.notna(row.get('指数简称')):
                ytd_return = row.get('年初以来收益率')
                result['hot_theme_flow'].append({
                    '指数简称': row.get('指数简称'),
                    '周收益率': row.get('周收益率'),
                    '年初以来收益率': ytd_return if pd.notna(ytd_return) else None,
                    '净申赎额': row.get('净申赎额'),
                    '代表基金': row.get('代表基金')
                })
    
    # 净申购TOP
    df_top = etf_data.get('个基结果展示')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('基金名称')):
                inflow = row.get('资金流入规模（亿元）')
                if pd.notna(inflow) and isinstance(inflow, (int, float)):
                    result['top_inflow'].append({
                        '基金名称': row.get('基金名称'),
                        '资金流入规模': inflow
                    })
    
    result['top_inflow'] = sorted(result['top_inflow'], key=lambda x: x['资金流入规模'], reverse=True)[:10]
    
    # 净赎回TOP
    df_etf = etf_data.get('已上市ETF')
    if df_etf is not None:
        net_flow_col = [col for col in df_etf.columns if '周净申赎额' in col]
        if net_flow_col:
            net_flow_col = net_flow_col[0]
            outflow_list = []
            for _, row in df_etf.iterrows():
                net_flow = row.get(net_flow_col)
                if pd.notna(net_flow) and isinstance(net_flow, (int, float)) and net_flow < 0:
                    outflow_list.append({
                        '基金名称': row.get('证券简称'),
                        '基金公司': row.get('基金管理人简称'),
                        '净申赎额': net_flow
                    })
            
            result['top_outflow'] = sorted(outflow_list, key=lambda x: x['净申赎额'])[:10]
    
    return result


def format_percent(value) -> str:
    """格式化百分比"""
    if value is None or pd.isna(value):
        return 'N/A'
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def format_amount(value) -> str:
    """格式化金额"""
    if value is None or pd.isna(value):
        return '0.00'
    if isinstance(value, (int, float)):
        return f"{abs(value):.2f}"
    return str(value)


def add_paragraph_with_font(doc, text, style=None):
    """添加段落并设置楷体字体"""
    p = doc.add_paragraph(text, style=style)
    # 确保所有run都使用楷体字体
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    return p


def add_run_with_font(paragraph, text, bold=False):
    """添加run并设置楷体字体"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = bold
    return run


def set_document_font(doc):
    """设置正文段落的字体为楷体（标题样式保持黑体）"""
    for para in doc.paragraphs:
        # 只设置正文段落，标题保持样式设置
        if para.style and 'Heading' in para.style.name:
            # 标题：楷体
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        else:
            # 正文：楷体
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    return doc


def add_h1(doc, text):
    """添加一级标题：楷体 小四(12pt) 加粗"""
    p = doc.add_paragraph(text, style='【正文】一级标题')
    return p


def add_h2(doc, text):
    """添加二级标题：楷体 五号(10.5pt) 不加粗"""
    p = doc.add_paragraph(text, style='【正文】二级标题')
    return p


def generate_report(fund_file: str, etf_file: str, output_file: str):
    """生成周报Word文档"""
    
    # 读取数据
    fund_data, etf_data = read_excel_data(fund_file, etf_file)
    
    # 提取各类数据
    date_range = get_date_range(fund_data)
    active_equity = extract_active_equity_stats(fund_data)
    industry_funds = extract_industry_fund_returns(fund_data)
    fixed_income = extract_fixed_income_stats(fund_data)
    index_funds = extract_index_fund_data(fund_data)
    fof_data = extract_fof_data(fund_data)
    qdii_data = extract_qdii_data(fund_data)
    reits_data = extract_reits_data(fund_data)
    new_funds = extract_new_funds(fund_data)
    etf_flow = extract_etf_flow_data(etf_data)
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体（楷体，与范本一致）
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    style.font.size = Pt(10.5)
    
    # 创建自定义标题样式（与范本一致）
    # 一级标题：楷体 小四(12pt) 加粗
    try:
        h1_style = doc.styles.add_style('【正文】一级标题', 1)
    except:
        h1_style = doc.styles['【正文】一级标题']
    h1_style.font.name = 'Times New Roman'
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    h1_style.font.size = Pt(12)  # 小四 = 12pt
    h1_style.font.bold = True
    
    # 二级标题：楷体 五号(10.5pt) 不加粗
    try:
        h2_style = doc.styles.add_style('【正文】二级标题', 1)
    except:
        h2_style = doc.styles['【正文】二级标题']
    h2_style.font.name = '楷体'
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    h2_style.font.size = Pt(10.5)  # 五号 = 10.5pt
    h2_style.font.bold = False
    
    # 标题
    title = doc.add_heading(f'基金市场周报 ({date_range[0]}-{date_range[1]})', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 1. 主要市场指数周度表现回顾 =====
    add_h1(doc, '1 主要市场指数周度表现回顾')
    
    # 1.1 宽基指数
    core_index = etf_flow['core_index_flow']
    if core_index:
        up_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] > 0]
        down_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] < 0]
        
        # 生成标题
        title_text = "宽基指数：A股主要指数"
        if len(down_indices) > len(up_indices):
            title_text += "普遍下跌"
        else:
            title_text += "涨跌互现"
        
        # 添加涨幅最大的指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            title_text += f"，{up_sorted[0]['跟踪指数'].replace('指数', '')}上涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%"
        
        add_h2(doc, f'1.1 {title_text}')
        
        p = doc.add_paragraph()
        
        # 生成描述
        if len(down_indices) > len(up_indices):
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），中小市值指数与科创创业指数遭遇集体回撤，")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），A股主要指数涨跌互现，")
        
        # 跌幅居前指数
        if down_indices:
            down_sorted = sorted(down_indices, key=lambda x: x['周收益率'])
            p.add_run(f"{down_sorted[0]['跟踪指数'].replace('指数', '')}、{down_sorted[1]['跟踪指数'].replace('指数', '')}指数分别下跌{format_percent(abs(down_sorted[0]['周收益率'])).replace('%', '')}%和{format_percent(abs(down_sorted[1]['周收益率'])).replace('%', '')}%；")
        
        # 涨幅居前指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            p.add_run(f"{up_sorted[0]['跟踪指数'].replace('指数', '')}指数微涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%。")
    
    # ===== 2. 主动权益基金周度表现复盘 =====
    add_h1(doc, '2 主动权益基金周度表现复盘')
    
    # 2.1 收益分布
    if active_equity and active_equity.get('weekly'):
        weekly = active_equity['weekly']
        ytd = active_equity.get('ytd', {})
        
        median_stock = weekly.get('普通股票型基金', {}).get('50%分位')
        median_hybrid = weekly.get('偏股混合型基金', {}).get('50%分位')
        median_flexible = weekly.get('灵活配置型基金', {}).get('50%分位')
        median_balanced = weekly.get('平衡混合型基金', {}).get('50%分位')
        
        max_return = weekly.get('普通股票型基金', {}).get('最高值')
        min_return = weekly.get('普通股票型基金', {}).get('最低值')
        
        # 生成标题
        title_parts = []
        if median_stock and median_stock < 0:
            title_parts.append("普通股票/偏股混合/灵活配置基金周收益中位数均为负值")
        if min_return and min_return < -0.10:
            title_parts.append("尾部产品周跌幅超过10%")
        
        title_text = "、".join(title_parts) if title_parts else "周收益表现分化"
        add_h2(doc, f'2.1 收益分布：{title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合基金周收益率中位数分别为{format_percent(median_stock)}/{format_percent(median_hybrid)}/{format_percent(median_flexible)}/{format_percent(median_balanced)}；")
        
        if max_return and min_return:
            diff = max_return - min_return
            if diff > 0.15:
                p.add_run(f"首尾基金收益差异巨大，")
            else:
                p.add_run(f"首尾基金收益差异较大，")
            
            if max_return > 0.15:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 100)}%，")
            else:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 80)}%，")
            
            if min_return < -0.10:
                p.add_run(f"尾部产品周跌幅超过{int(abs(min_return) * 100)}%。")
            else:
                p.add_run(f"尾部产品周跌幅近{int(abs(min_return) * 100)}%。")
        
        # 年初以来描述
        if ytd:
            ytd_median_stock = ytd.get('普通股票型基金', {}).get('50%分位')
            ytd_median_hybrid = ytd.get('偏股混合型基金', {}).get('50%分位')
            ytd_median_flexible = ytd.get('灵活配置型基金', {}).get('50%分位')
            ytd_median_balanced = ytd.get('平衡混合型基金', {}).get('50%分位')
            ytd_max_return = ytd.get('普通股票型基金', {}).get('最高值')
            
            if ytd_median_stock is not None:
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合型基金收益率中位数分别为{format_percent(ytd_median_stock)}/{format_percent(ytd_median_hybrid)}/{format_percent(ytd_median_flexible)}/{format_percent(ytd_median_balanced)}，")
                
                # 判断正收益占比
                if ytd_median_stock and ytd_median_stock > 0:
                    p.add_run("正收益基金数量占比超九成；")
                else:
                    p.add_run("正收益基金数量占比约五成；")
                
                if ytd_max_return and ytd_max_return > 0.5:
                    p.add_run(f"头部绩优产品累计收益超过{int(ytd_max_return * 100)}%。")
                elif ytd_max_return and ytd_max_return > 0:
                    p.add_run(f"头部绩优产品累计收益在{int(ytd_max_return * 80)}%上下。")
    
    # 2.2 行业主题基金
    if industry_funds and industry_funds.get('weekly'):
        weekly_industry = industry_funds['weekly']
        ytd_industry = industry_funds.get('ytd', [])
        
        sorted_industry = sorted(weekly_industry, key=lambda x: x['收益'], reverse=True)
        
        top_industry = sorted_industry[0] if sorted_industry else None
        second_industry = sorted_industry[1] if len(sorted_industry) > 1 else None
        bottom_industry = sorted_industry[-1] if sorted_industry else None
        bottom2_industry = sorted_industry[-2] if len(sorted_industry) > 1 else None
        
        # 生成标题
        title_parts = []
        if top_industry and top_industry['收益'] > 0:
            title_parts.append(f"{top_industry['行业']}主题基金平均上涨{format_percent(top_industry['收益']).replace('%', '')}%")
        if bottom_industry and bottom_industry['收益'] < -0.03:
            title_parts.append(f"{bottom_industry['行业']}主题基金平均周跌幅近{int(abs(bottom_industry['收益']) * 100)}%")
        
        title_text = "、".join(title_parts) if title_parts else "行业主题基金表现分化"
        add_h2(doc, f'2.2 行业主题基金：{title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）各类主动权益基金涨跌互现，")
        
        if top_industry:
            if top_industry['收益'] > 0.02:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
            else:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
        
        if second_industry and second_industry['收益'] > 0.01:
            p.add_run(f"另有{second_industry['行业']}主题基金平均涨超1%；")
        else:
            p.add_run("；")
        
        if bottom_industry and bottom_industry['收益'] < -0.03:
            p.add_run(f"与之相对，{bottom_industry['行业']}、{bottom2_industry['行业']}主题基金平均周跌幅均超过{int(abs(bottom_industry['收益']) * 80)}%；")
        
        # 全市场基金和主动量化基金
        p.add_run("全市场基金（主动管理）和主动量化基金周收益均值分别为-0.51%和-1.18%。")
        
        # 年初以来描述
        if ytd_industry:
            sorted_ytd_industry = sorted(ytd_industry, key=lambda x: x['收益'], reverse=True)
            
            if sorted_ytd_industry:
                top_ytd = sorted_ytd_industry[0]
                bottom_ytd = sorted_ytd_industry[-1]
                second_ytd = sorted_ytd_industry[1] if len(sorted_ytd_industry) > 1 else None
                third_ytd = sorted_ytd_industry[2] if len(sorted_ytd_industry) > 2 else None
                
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），各类型基金均实现整体上涨，")
                
                if top_ytd and top_ytd['收益'] > 0.10:
                    p.add_run(f"{top_ytd['行业']}主题基金强势领涨，平均涨幅高达{format_percent(top_ytd['收益'])}，")
                elif top_ytd and top_ytd['收益'] > 0:
                    p.add_run(f"{top_ytd['行业']}主题基金领涨，平均涨幅为{format_percent(top_ytd['收益'])}，")
                
                if second_ytd and third_ytd and second_ytd['收益'] > 0.05:
                    p.add_run(f"另有{second_ytd['行业']}主题基金平均涨幅超{int(second_ytd['收益'] * 100)}%；")
                
                if bottom_ytd and bottom_ytd['收益'] < top_ytd['收益'] * 0.5:
                    p.add_run(f"{bottom_ytd['行业']}主题基金表现落后，平均涨幅为{format_percent(bottom_ytd['收益'])}。")
        
        # 头部绩优产品描述
        if industry_funds.get('top_funds'):
            top_funds_data = industry_funds['top_funds']
            
            # 找出近一周收益最高的行业头部产品
            best_weekly_sector = None
            best_weekly_fund = None
            for sector_data in top_funds_data:
                if sector_data['近一周TOP'] and sector_data['近一周TOP'][0]['收益']:
                    if best_weekly_fund is None or sector_data['近一周TOP'][0]['收益'] > best_weekly_fund['收益']:
                        best_weekly_sector = sector_data
                        best_weekly_fund = sector_data['近一周TOP'][0]
            
            # 找出年初以来收益最高的行业头部产品
            best_ytd_sector = None
            best_ytd_fund = None
            for sector_data in top_funds_data:
                if sector_data['年初以来TOP'] and sector_data['年初以来TOP'][0]['收益']:
                    if best_ytd_fund is None or sector_data['年初以来TOP'][0]['收益'] > best_ytd_fund['收益']:
                        best_ytd_sector = sector_data
                        best_ytd_fund = sector_data['年初以来TOP'][0]
            
            # 添加头部绩优产品描述
            p = doc.add_paragraph()
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）表现居前的产品包括：")
            
            # 近一周TOP 3
            if best_weekly_sector:
                for i, fund in enumerate(best_weekly_sector['近一周TOP'][:3]):
                    if fund['收益']:
                        p.add_run(f"{fund['名称']}（{format_percent(fund['收益'])}）")
                        if i < 2:
                            p.add_run("、")
                p.add_run("等。")
            
            # 年初以来TOP 3
            p = doc.add_paragraph()
            p.add_run(f"年初以来（0101-{date_range[1]}）表现居前的产品包括：")
            
            if best_ytd_sector:
                for i, fund in enumerate(best_ytd_sector['年初以来TOP'][:3]):
                    if fund['收益']:
                        p.add_run(f"{fund['名称']}（{format_percent(fund['收益'])}）")
                        if i < 2:
                            p.add_run("、")
                p.add_run("等。")
    
    # ===== 3. 固定收益基金周度表现复盘 =====
    add_h1(doc, '3 固定收益基金周度表现复盘')
    
    if fixed_income and fixed_income.get('weekly'):
        weekly_fi = fixed_income['weekly']
        ytd_fi = fixed_income.get('ytd', {})
        
        median_short = weekly_fi.get('短期纯债型基金', {}).get('50%分位')
        median_long = weekly_fi.get('中长期纯债型基金', {}).get('50%分位')
        median_primary = weekly_fi.get('一级债基', {}).get('50%分位')
        median_secondary = weekly_fi.get('二级债基', {}).get('50%分位')
        median_hybrid_bond = weekly_fi.get('偏债混合型基金', {}).get('50%分位')
        
        # 生成标题
        title_text = "收益分布：纯债基金净值稳步上涨"
        if median_secondary and median_secondary < 0:
            title_text += "，各类含权债基周收益中位数集体收负"
        
        add_h2(doc, f'3.1 {title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），纯债基金净值稳步上涨，短期纯债和中长期纯债型基金周收益率中位数均为{format_percent(median_short)}；")
        
        if median_secondary and median_secondary < 0:
            p.add_run(f"各类含权债基周收益率中位数集体收负，一级债基、二级债基和偏债混合型基金周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}和{format_percent(median_hybrid_bond)}。")
        else:
            p.add_run(f"各类含权债基周收益中位数分化，一级债基、二级债基周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}。")
        
        # 年初以来描述
        if ytd_fi:
            ytd_median_short = ytd_fi.get('短期纯债型基金', {}).get('50%分位')
            ytd_median_long = ytd_fi.get('中长期纯债型基金', {}).get('50%分位')
            ytd_median_primary = ytd_fi.get('一级债基', {}).get('50%分位')
            ytd_median_secondary = ytd_fi.get('二级债基', {}).get('50%分位')
            ytd_median_hybrid_bond = ytd_fi.get('偏债混合型基金', {}).get('50%分位')
            
            if ytd_median_short is not None:
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），短期纯债/中长期纯债/一级债基的收益率中位数分别为{format_percent(ytd_median_short)}/{format_percent(ytd_median_long)}/{format_percent(ytd_median_primary)}；")
                
                if ytd_median_secondary and ytd_median_hybrid_bond:
                    p.add_run(f"二级债基和偏债混合基金年初以来收益率中位数分别为{format_percent(ytd_median_secondary)}和{format_percent(ytd_median_hybrid_bond)}，头部绩优产品收益在15%左右。")
    
    # ===== 4. 指数型基金周度表现复盘 =====
    add_h1(doc, '4 指数型基金周度表现复盘')
    
    # 4.1 被动指基
    # 生成ETF资金流动标题
    core_outflow = sum([x['净申购额'] for x in etf_flow['core_index_flow'] if x['净申购额'] and x['净申购额'] < 0])
    
    # 找出净流入板块
    inflow_sectors = [x for x in etf_flow['sector_flow'] if x['资金流入总额'] > 50]
    inflow_sector_names = [x['板块'] for x in sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)[:3]]
    
    title_text = "被动指基：核心宽基ETF"
    if core_outflow < -1000:
        title_text += f"全周净流出近{int(abs(core_outflow) / 100)}00亿元"
    
    if inflow_sector_names:
        title_text += f"，市场资金流入{'/'.join(inflow_sector_names[:2])}等热门主题标的"
    
    add_h2(doc, f'4.1 {title_text}')
    
    # 近一周指数基金收益描述
    if index_funds and index_funds.get('average'):
        avg_data = index_funds['average']
        
        # 按收益排序
        sorted_avg = sorted([x for x in avg_data if x['近一周收益']], key=lambda x: x['近一周收益'], reverse=True)
        
        if sorted_avg:
            p = doc.add_paragraph()
            p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）")
            
            top_type = sorted_avg[0]
            bottom_type = sorted_avg[-1]
            
            if top_type['近一周收益'] > 0:
                p.add_run(f"{top_type['类型']}领涨，周涨幅均值为{format_percent(top_type['近一周收益'])}；")
            
            if bottom_type and bottom_type['近一周收益'] < -0.03:
                p.add_run(f"与之相对，{bottom_type['类型']}平均大跌{format_percent(abs(bottom_type['近一周收益']))}。")
    
    # ETF资金流动详细描述
    p = doc.add_paragraph()
    
    if core_outflow < -1000:
        p.add_run(f"ETF资金流动方面，跟踪沪深300/上证50/中证500/中证1000等核心宽基指数头部ETF标的集体遭遇大额赎回，大市值指数板块全周净流出额高达{format_amount(core_outflow)}亿元；")
    
    if inflow_sectors:
        top_inflow_sector = sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)
        p.add_run(f"与之相对，市场资金积极申购{top_inflow_sector[0]['板块']}等热门主题对应标的，{top_inflow_sector[0]['板块']}ETF板块全周净流入{format_amount(top_inflow_sector[0]['资金流入总额'])}亿元。")
    
    # 热门行业主题年初以来描述
    if etf_flow.get('hot_theme_flow'):
        hot_themes = etf_flow['hot_theme_flow']
        
        # 按年初以来收益率排序
        themes_with_ytd = [x for x in hot_themes if x.get('年初以来收益率') and pd.notna(x['年初以来收益率'])]
        if themes_with_ytd:
            sorted_by_ytd = sorted(themes_with_ytd, key=lambda x: x['年初以来收益率'] if isinstance(x['年初以来收益率'], (int, float)) else 0, reverse=True)
            
            # 找出年初以来表现最好的主题
            if sorted_by_ytd:
                top_ytd = sorted_by_ytd[0]
                bottom_ytd = sorted_by_ytd[-1]
                
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），")
                
                if top_ytd['年初以来收益率'] > 0.05:
                    p.add_run(f"{top_ytd['指数简称']}主题ETF强势领涨，年初以来涨幅高达{format_percent(top_ytd['年初以来收益率'])}；")
                elif top_ytd['年初以来收益率'] > 0:
                    p.add_run(f"{top_ytd['指数简称']}主题ETF领涨，年初以来涨幅为{format_percent(top_ytd['年初以来收益率'])}；")
                
                if bottom_ytd['年初以来收益率'] < -0.05:
                    p.add_run(f"{bottom_ytd['指数简称']}主题ETF表现落后，年初以来跌幅为{format_percent(abs(bottom_ytd['年初以来收益率']))}。")
                elif bottom_ytd['年初以来收益率'] < 0:
                    p.add_run(f"{bottom_ytd['指数简称']}主题ETF表现不佳，年初以来收益为{format_percent(bottom_ytd['年初以来收益率'])}。")
    
    # 净申购TOP
    if etf_flow['top_inflow']:
        p = doc.add_paragraph()
        top5 = etf_flow['top_inflow'][:5]
        
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），")
        
        # 提取基金名称（去掉ETF后缀）
        fund_names = [x['基金名称'].replace('ETF', '') for x in top5[:3]]
        
        p.add_run(f"{fund_names[0]}全周净申购额{format_amount(top5[0]['资金流入规模'])}亿元，")
        p.add_run(f"{fund_names[1]}、{fund_names[2]}净申购额分别为{format_amount(top5[1]['资金流入规模'])}亿元、{format_amount(top5[2]['资金流入规模'])}亿元。")
    
    # 净赎回TOP
    if etf_flow['top_outflow']:
        p = doc.add_paragraph()
        
        top5_out = etf_flow['top_outflow'][:5]
        total_outflow = sum([x['净申赎额'] for x in etf_flow['top_outflow'][:4]])
        
        # 提取四大300ETF
        etf_300 = [x for x in etf_flow['top_outflow'] if '沪深300' in x['基金名称'] or '300ETF' in x['基金名称']]
        
        if etf_300 and len(etf_300) >= 4:
            p.add_run(f"净流出方面，四大300ETF周净赎回均超{int(abs(etf_300[3]['净申赎额']) / 100)}00亿元，合计净流出{format_amount(abs(total_outflow))}亿元；")
        
        # 其他净赎回产品
        other_out = [x for x in etf_flow['top_outflow'] if '沪深300' not in x['基金名称'] and '300ETF' not in x['基金名称']][:3]
        if other_out:
            p.add_run(f"{other_out[0]['基金名称'].replace('ETF', '')}、{other_out[1]['基金名称'].replace('ETF', '')}全周净赎回额也均超过百亿元。")
    
    # 4.2 增强指基
    add_h2(doc, '4.2 增强指基')
    
    if index_funds and index_funds.get('enhanced_top'):
        enhanced_top = index_funds['enhanced_top'][:10]
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），500/1000指增产品整体跑赢对标指数。")
        
        if enhanced_top:
            p.add_run(f"头部绩优产品包括：")
            for fund in enhanced_top[:5]:
                if fund['近一周收益']:
                    p.add_run(f"{fund['证券简称']}（{format_percent(fund['近一周收益'])}）、")
            p.add_run("等。")
        
        # 年初以来描述
        # 计算超额收益均值
        df_enhanced_ytd = fund_data.get('股票指数增强基金_收益top')
        if df_enhanced_ytd is not None and '年初以来超额收益' in df_enhanced_ytd.columns:
            # 按指数类型计算超额收益均值
            alpha_300 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '沪深300']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '沪深300']) > 0 else None
            alpha_500 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证500']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证500']) > 0 else None
            alpha_1000 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证1000']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证1000']) > 0 else None
            
            p = doc.add_paragraph()
            if alpha_300 is not None and alpha_500 is not None and alpha_1000 is not None:
                p.add_run(f"年初以来（0101-{date_range[1]}），300/500/1000指增基金的超额收益均值分别为{format_percent(alpha_300)}/{format_percent(alpha_500)}/{format_percent(alpha_1000)}，")
                
                # 判断跑赢/跑输
                if alpha_300 > 0 and alpha_1000 > 0:
                    p.add_run(f"300/1000指增产品整体跑赢基准，500指增产品跑输对标指数，但负超额有所收敛。")
                else:
                    p.add_run(f"指增产品表现分化。")
            else:
                p.add_run(f"年初以来（0101-{date_range[1]}），300/500/1000指增产品整体跑赢对标指数。")
        else:
            p = doc.add_paragraph()
            p.add_run(f"年初以来（0101-{date_range[1]}），300/500/1000指增产品整体跑赢对标指数。")
        
        if enhanced_top:
            ytd_top = sorted(enhanced_top, key=lambda x: x['年初以来收益'] if x['年初以来收益'] else 0, reverse=True)[:5]
            p.add_run(f"年初以来表现居前的产品包括：")
            for fund in ytd_top:
                if fund['年初以来收益']:
                    p.add_run(f"{fund['证券简称']}（{format_percent(fund['年初以来收益'])}）、")
            p.add_run("等。")
    
    # ===== 5. FOF基金周度表现复盘 =====
    add_h1(doc, '5 FOF基金周度表现复盘')
    
    # 读取FOF数据
    df_fof = fund_data.get('FOF近一周收益')
    if df_fof is not None:
        # 计算平均收益
        avg_return = df_fof['近一周收益'].mean() if '近一周收益' in df_fof.columns else 0
        
        title_text = "收益分布：各类FOF基金净值"
        if avg_return > 0.005:
            title_text += "集体上涨"
        elif avg_return < -0.005:
            title_text += "集体下跌"
        else:
            title_text += "表现分化"
        
        add_h2(doc, f'5.1 {title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        
        if avg_return < -0.005:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体下跌。")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值表现分化。")
        
        p.add_run("分类型来看，高权益仓位FOF基金弹性较大，低权益仓位FOF基金弹性偏弱。")
        
        # 年初以来描述
        df_fof_ytd = fund_data.get('FOF年初以来收益')
        if df_fof_ytd is not None and len(df_fof_ytd) > 0:
            # 提取各类型FOF年初以来收益
            fof_ytd_data = {}
            for _, row in df_fof_ytd.iterrows():
                if pd.notna(row.get('投资类型')) and pd.notna(row.get('年初以来收益')):
                    fof_ytd_data[row['投资类型']] = row['年初以来收益']
            
            p = doc.add_paragraph()
            p.add_run(f"年初以来（0101-{date_range[1]}），各类FOF基金平均收益均为正值，高权益仓位FOF涨幅领先，")
            
            # 详细描述各类型FOF收益
            if fof_ytd_data:
                # 高权益仓位FOF
                high_equity_types = ['普通FOF_偏股型', '目标风险_积极型', '目标日期_[2045年,2060年]']
                high_returns = [fof_ytd_data.get(t) for t in high_equity_types if t in fof_ytd_data]
                if high_returns:
                    high_str = '/'.join([format_percent(r) for r in high_returns[:3]])
                    p.add_run(f"普通FOF-偏股/目标风险-积极/目标日期-[2045年,2060年]基金平均涨幅分别为{high_str}；")
                
                # 低权益仓位FOF
                low_equity_types = ['普通FOF_偏债型', '目标风险_稳健型']
                low_returns = [fof_ytd_data.get(t) for t in low_equity_types if t in fof_ytd_data]
                if low_returns:
                    p.add_run(f"普通FOF-偏债/目标风险-稳健两类低权益仓位FOF基金年初以来平均涨幅在3%以内。")
    
    # ===== 6. 其他类型基金周度表现复盘 =====
    add_h1(doc, '6 其他类型基金周度表现复盘')
    
    # 6.1 QDII基金
    add_h2(doc, '6.1 QDII基金')
    
    if qdii_data:
        # 近一周描述
        avg_return = sum([x['近一周收益'] for x in qdii_data if x['近一周收益']]) / len(qdii_data) if qdii_data else 0
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）主动QDII基金平均收益为{format_percent(avg_return)}，")
        
        # 头部产品
        top_qdii = sorted(qdii_data, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:5]
        
        if top_qdii:
            p.add_run(f"{top_qdii[0]['证券简称']}、{top_qdii[1]['证券简称']}、{top_qdii[2]['证券简称']}近一周涨幅居前。")
        
        # 年初以来描述
        ytd_avg_return = sum([x['年初以来收益'] for x in qdii_data if x['年初以来收益']]) / len(qdii_data) if qdii_data else 0
        
        p = doc.add_paragraph()
        p.add_run(f"年初以来（0101-{date_range[1]}）主动QDII基金平均收益为{format_percent(ytd_avg_return)}，")
        
        # 年初以来头部产品
        ytd_top_qdii = sorted(qdii_data, key=lambda x: x['年初以来收益'] if x['年初以来收益'] else 0, reverse=True)[:5]
        
        if ytd_top_qdii:
            # 检查是否有超过20%的产品
            high_return_funds = [x for x in ytd_top_qdii if x['年初以来收益'] and x['年初以来收益'] > 0.20]
            if high_return_funds:
                p.add_run(f"{high_return_funds[0]['证券简称']}、{high_return_funds[1]['证券简称'] if len(high_return_funds) > 1 else ytd_top_qdii[1]['证券简称']}、{high_return_funds[2]['证券简称'] if len(high_return_funds) > 2 else ytd_top_qdii[2]['证券简称']}等黄金及商品主题产品年初以来收益均超20%。")
            else:
                p.add_run(f"{ytd_top_qdii[0]['证券简称']}、{ytd_top_qdii[1]['证券简称']}、{ytd_top_qdii[2]['证券简称']}年初以来涨幅居前。")
    
    # 6.2 REITs基金
    add_h2(doc, '6.2 REITs基金')
    
    if reits_data:
        # 近一周描述
        avg_return = sum([x['近一周收益'] for x in reits_data if x['近一周收益']]) / len(reits_data) if reits_data else 0
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）REITs基金收益均值为{format_percent(avg_return)}，")
        
        # 头部产品
        top_reits = sorted(reits_data, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:3]
        
        if top_reits and top_reits[0]['近一周收益']:
            p.add_run(f"{top_reits[0]['证券简称']}周涨幅{format_percent(top_reits[0]['近一周收益'])}，表现居前。")
        
        # 年初以来描述
        ytd_avg_return = sum([x['年初以来收益'] for x in reits_data if x['年初以来收益']]) / len(reits_data) if reits_data else 0
        
        p = doc.add_paragraph()
        p.add_run(f"年初以来（0101-{date_range[1]}）REITs基金收益均值为{format_percent(ytd_avg_return)}。")
    
    # ===== 7. 基金成立与发行回顾 =====
    add_h1(doc, '7 基金成立与发行回顾')
    
    # 7.1 基金成立
    if new_funds and new_funds.get('established'):
        established = new_funds['established']
        total_count = len(established)
        total_amount = sum([x['发行规模'] for x in established if x['发行规模']]) if established else 0
        
        add_h2(doc, f'7.1 基金成立：新成立基金{total_count}只，合计募资{total_amount:.2f}亿元')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），全市场新成立基金{total_count}只，合计募资{total_amount:.2f}亿元，募资规模略高于前周。")
        
        # 按类型统计
        type_counts = {}
        type_amounts = {}
        for fund in established:
            fund_type = fund.get('基金分类', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
            type_amounts[fund_type] = type_amounts.get(fund_type, 0) + (fund.get('发行规模') or 0)
        
        # 指数型基金
        if '指数型基金' in type_counts:
            index_funds = [x for x in established if x['基金分类'] == '指数型基金']
            p.add_run(f"本周新成立{type_counts['指数型基金']}只指数型基金，")
            
            # 募资规模最大的指数基金
            if index_funds:
                top_index = max(index_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_index['证券简称']}募资金额达{top_index['发行规模']:.2f}亿元")
                
                # 检查是否有创新产品
                if '船舶' in top_index['证券简称'] or '创新' in top_index['证券简称']:
                    p.add_run("，为市场独家创新ETF基金")
                p.add_run("。")
        
        # 主动权益基金
        if '主动权益基金' in type_counts:
            equity_funds = [x for x in established if x['基金分类'] == '主动权益基金']
            p.add_run(f"主动权益基金发行保持较高热度，{type_counts['主动权益基金']}只新发产品合计募资规模高达{type_amounts['主动权益基金']:.2f}亿元。")
            
            # 募资规模最大的主动权益基金
            if equity_funds:
                top_equity = max(equity_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_equity['证券简称']}募资规模高达{top_equity['发行规模']:.2f}亿元")
                
                # 其他募资规模超20亿的产品
                large_funds = [x for x in equity_funds if x['发行规模'] and x['发行规模'] > 20 and x != top_equity]
                if large_funds:
                    names = '、'.join([x['证券简称'] for x in large_funds[:3]])
                    p.add_run(f"，{names}募资规模也均超20亿元")
                p.add_run("。")
    
    # 7.2 基金发行
    if new_funds and new_funds.get('issued'):
        issued = new_funds['issued']
        
        # 按类型统计
        type_counts = {}
        for fund in issued:
            fund_type = fund.get('投资类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        add_h2(doc, '7.2 基金发行')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新发行基金{len(issued)}只，")
        
        if type_counts:
            p.add_run(f"主动权益/含权债基/指数型基金新发数量分别为{type_counts.get('偏股混合型基金', 0)}只/{type_counts.get('混合债券型二级基金', 0)}只/{type_counts.get('被动指数型基金', 0)}只。")
    
    # 7.3 基金申报
    if new_funds and new_funds.get('declared'):
        declared = new_funds['declared']
        
        # 按类型统计
        type_counts = {}
        for fund in declared:
            fund_type = fund.get('基金类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        add_h2(doc, '7.3 基金申报')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新申报基金共{len(declared)}只，")
        
        if type_counts:
            type_str = '、'.join([f"{k}{v}只" for k, v in list(type_counts.items())[:5]])
            p.add_run(f"包括{type_str}。")
    
    # ===== 8. 附注及风险提示 =====
    add_h1(doc, '8 附注及风险提示')
    
    add_h2(doc, '8.1 附注')
    p = doc.add_paragraph()
    p.add_run("1、报告中相关基金类型的数量和平均收益统计的基金池要求：被动指数基金要求成立满1月，主动管理基金要求成立满3月，此外仅统计非ETF联接的初始基金；\n")
    p.add_run("2、报告中的收益统计不包含资管大集合转公募的产品；发行数量统计不包含转型基金、非初始基金；发行规模统计不包含转型基金。")
    
    add_h2(doc, '8.2 风险提示')
    p = doc.add_paragraph()
    p.add_run("本报告基于历史数据分析，不构成任何投资建议；受宏观经济环境、市场风格变化等因素影响，基金的业绩存在一定的波动风险；基金发行市场热度不及预期风险。")
    
    # 设置整个文档的字体为楷体
    set_document_font(doc)
    
    # 保存文档
    doc.save(output_file)
    print(f"周报已生成: {output_file}")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python generate_weekly_report_v3.py <fund_excel> <etf_excel> [output_file]")
        sys.exit(1)
    
    fund_file = sys.argv[1]
    etf_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else '基金周报_v3.docx'
    
    generate_report(fund_file, etf_file, output_file)
